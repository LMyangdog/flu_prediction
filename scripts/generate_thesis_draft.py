"""
Generate the undergraduate thesis draft Word document for the flu_prediction project.

The script creates:
  - results/figures/thesis_system_architecture.png
  - results/figures/thesis_data_pipeline.png
  - 基于深度学习和多元数据的流感爆发趋势预测_毕业论文初稿.docx
"""

from __future__ import annotations

import csv
import json
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "results" / "figures"
REPORT_DIR = ROOT / "results" / "reports"
OUT_DOCX = ROOT / "基于深度学习和多元数据的流感爆发趋势预测_毕业论文初稿.docx"


TITLE = "基于深度学习和多元数据的流感爆发趋势预测"
STUDENT = "张宇鑫"
STUDENT_ID = "22090032057"
SCHOOL = "信息科学与工程学部"
MAJOR = "计算机科学与技术2022级"
ADVISOR = "刘艳艳"


def load_json(path: Path) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_csv(path: Path) -> list[dict]:
    with open(path, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def font_path() -> str:
    for path in [r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simsun.ttc"]:
        if os.path.exists(path):
            return path
    return ""


def draw_flow_diagram(path: Path, title: str, nodes: list[str], colors: list[str]) -> None:
    w, h = 1800, 760
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 54) if fp else ImageFont.load_default()
    node_font = ImageFont.truetype(fp, 34) if fp else ImageFont.load_default()
    small_font = ImageFont.truetype(fp, 26) if fp else ImageFont.load_default()

    draw.text((w / 2, 60), title, fill="#1F2937", font=title_font, anchor="mm")
    box_w, box_h = 300, 140
    gap = (w - 2 * 120 - len(nodes) * box_w) / (len(nodes) - 1)
    y = 270
    for i, node in enumerate(nodes):
        x = 120 + i * (box_w + gap)
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=28, fill=colors[i], outline="#1F2937", width=3)
        lines = node.split("\n")
        for j, line in enumerate(lines):
            draw.text((x + box_w / 2, y + 50 + j * 42), line, fill="white", font=node_font, anchor="mm")
        if i < len(nodes) - 1:
            x1 = x + box_w + 15
            x2 = x + box_w + gap - 15
            yy = y + box_h / 2
            draw.line((x1, yy, x2, yy), fill="#374151", width=5)
            draw.polygon([(x2, yy), (x2 - 22, yy - 14), (x2 - 22, yy + 14)], fill="#374151")
    draw.text((w / 2, 610), "说明：所有数据按周粒度对齐，训练/验证/测试按时间顺序严格划分，避免未来信息泄漏。", fill="#4B5563", font=small_font, anchor="mm")
    img.save(path)


def create_thesis_figures() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    draw_flow_diagram(
        FIG_DIR / "thesis_data_pipeline.png",
        "多源数据处理与实验流程",
        ["原始数据\n采集留痕", "周粒度\n对齐融合", "特征工程\n构造", "严格时序\n切分", "模型训练\n与评估"],
        ["#2563EB", "#059669", "#D97706", "#7C3AED", "#DC2626"],
    )
    draw_flow_diagram(
        FIG_DIR / "thesis_system_architecture.png",
        "系统总体架构",
        ["数据层\nCSV/Manifest", "处理层\n清洗与特征", "模型层\niTransformer", "评估层\n指标与图表", "展示层\nStreamlit"],
        ["#0F766E", "#1D4ED8", "#9333EA", "#B45309", "#BE123C"],
    )


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def set_paragraph_format(p, first_line=True, align=None):
    fmt = p.paragraph_format
    fmt.line_spacing = Pt(24)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(24)
    if align is not None:
        p.alignment = align


def add_paragraph(doc: Document, text: str, first_line=True, align=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line, align=align)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_center(doc: Document, text: str, size=16, bold=True, font="黑体"):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, font, size, bold)
    return p


def add_heading(doc: Document, text: str, level: int):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", 16 if level == 1 else 14, True)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5)
    return p


def add_picture(doc: Document, path: Path, caption: str, width_inches=5.8):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    add_caption(doc, caption)
    for row in rows:
        label = str(row[0])
        details = []
        for idx in range(1, min(len(headers), len(row))):
            details.append(f"{headers[idx]}：{row[idx]}")
        text = f"{label}：" + "；".join(details) if details else label
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False)
        run = p.add_run(text)
        set_run_font(run, "宋体", 10.5)
    return None


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, "宋体", 9)


def add_toc(doc: Document):
    add_center(doc, "目   录", size=16, bold=True, font="黑体")
    entries = [
        (0, "摘 要", "2"),
        (0, "Abstract", "3"),
        (0, "1 绪论", "5"),
        (1, "1.1 研究背景", "5"),
        (1, "1.2 研究目的与意义", "5"),
        (1, "1.3 国内外研究现状", "5"),
        (1, "1.4 本文主要工作", "6"),
        (0, "2 数据来源与预处理", "6"),
        (1, "2.1 研究口径与数据边界", "6"),
        (1, "2.2 数据质量审计", "8"),
        (1, "2.3 特征工程", "9"),
        (0, "3 模型与系统设计", "10"),
        (1, "3.1 总体技术路线", "10"),
        (1, "3.2 iTransformer 模型", "11"),
        (1, "3.3 基准模型", "12"),
        (0, "4 实验设计与结果分析", "12"),
        (1, "4.1 实验设置", "12"),
        (1, "4.2 多模型对比", "12"),
        (1, "4.3 多步预测分 horizon 分析", "14"),
        (1, "4.4 消融实验", "14"),
        (1, "4.5 注意力可视化分析", "15"),
        (0, "5 Web 展示系统实现", "16"),
        (0, "6 不足与展望", "17"),
        (0, "7 结论", "17"),
        (0, "参考文献", "18"),
        (0, "致谢", "20"),
        (0, "附录", "21"),
    ]
    for level, title, page in entries:
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.left_indent = Pt(21 * level)
        fmt.line_spacing = Pt(22)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.tab_stops.add_tab_stop(Inches(5.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(f"{title}\t{page}")
        set_run_font(run, "宋体", 12 if level == 0 else 11, bold=(level == 0))
    doc.add_page_break()


def configure_doc(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = section.header.paragraphs[0].add_run(TITLE)
        set_run_font(hr, "宋体", 9)
        add_page_number(section)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)


def fmt(value: float | str, digits=3) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def main():
    create_thesis_figures()
    summary = load_json(REPORT_DIR / "experiment_summary.json")
    ablation_rows = load_csv(REPORT_DIR / "ablation_metrics.csv")
    quality = load_json(REPORT_DIR / "data_quality_report.json")
    final_audit = load_json(REPORT_DIR / "final_data_audit.json")

    doc = Document()
    configure_doc(doc)

    # Cover
    add_center(doc, "中国海洋大学", size=18, bold=True, font="黑体")
    add_center(doc, "本 科 毕 业 设 计", size=22, bold=True, font="黑体")
    for _ in range(4):
        doc.add_paragraph()
    add_center(doc, TITLE, size=18, bold=True, font="黑体")
    for _ in range(5):
        doc.add_paragraph()
    cover_rows = [
        ("学生姓名", STUDENT),
        ("学号", STUDENT_ID),
        ("学院（学部）", SCHOOL),
        ("专业年级", MAJOR),
        ("指导教师", ADVISOR),
        ("完成日期", "2026年5月"),
    ]
    for k, v in cover_rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(28)
        r = p.add_run(f"{k}：{v}")
        set_run_font(r, "宋体", 14)
    doc.add_page_break()

    # Chinese abstract
    add_center(doc, TITLE, size=16, bold=True, font="黑体")
    add_center(doc, "摘   要", size=16, bold=True, font="黑体")
    itr_metrics = summary["metrics"]["iTransformer"]
    abstract = (
        "流感等呼吸道传染病具有明显季节性和非线性波动特征，传统单变量统计模型难以充分利用气象、网络搜索等外部信息。"
        "本文围绕中国北方省份流感活动周度趋势预测任务，构建了一个融合国家流感中心周度监测数据、气象数据与百度指数的多源时间序列预测系统。"
        "本文采用中国国家流感中心《中国流感监测周报》中北方省份哨点医院 ILI% 作为目标变量（ili_rate），"
        "并将北方代表城市气象要素和搜索指数作为外生特征。系统实现了数据来源清单登记、数据质量审计、周粒度对齐、特征工程、"
        "严格时间顺序切分、iTransformer 建模、多模型对比和 Streamlit 可视化展示。实验结果表明，在当前真实周度数据版本下，"
        f"iTransformer 的整体 RMSE 为 {itr_metrics['RMSE']:.3f}，MAE 为 {itr_metrics['MAE']:.3f}，R2 为 {itr_metrics['R2']:.3f}。"
        "消融实验显示，流感历史特征仍是最稳定的信息来源，搜索指数和气象特征对部分指标有辅助作用。"
        "本文不再使用旧版平谷代理序列和 synthetic_extension 结果作为最终实验结论。"
    )
    add_paragraph(doc, abstract)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    r = p.add_run("关键词：")
    set_run_font(r, "黑体", 12, True)
    r = p.add_run("流感预测；多源数据融合；iTransformer；时间序列；百度指数")
    set_run_font(r, "宋体", 12)
    doc.add_page_break()

    # English abstract
    add_center(doc, "Influenza Outbreak Trend Prediction Based on Deep Learning and Multi-source Data", size=16, bold=True, font="Times New Roman")
    add_center(doc, "Abstract", size=16, bold=True, font="Times New Roman")
    eng = (
        "Influenza activity shows strong seasonality and nonlinear fluctuations, which makes early trend prediction difficult for traditional "
        "single-source forecasting models. This thesis develops a weekly influenza trend prediction system based on multi-source time series data, "
        "including official influenza surveillance records from the Chinese National Influenza Center, meteorological variables and Baidu search indices. "
        "The target variable is the weekly ILI percentage of northern China (ili_rate), while meteorological variables and search indices aggregated "
        "from representative northern cities are used as exogenous features. The system implements data source registration, data quality auditing, "
        "weekly alignment, feature engineering, leakage-free time-based splitting, iTransformer modeling, baseline comparison and Streamlit visualization. "
        f"Experimental results on the current real weekly dataset show that iTransformer obtains an RMSE of {itr_metrics['RMSE']:.3f}, "
        f"an MAE of {itr_metrics['MAE']:.3f} and an R2 of {itr_metrics['R2']:.3f}. "
        "The ablation study indicates that historical influenza activity remains the most stable source of predictive information, while search "
        "indices and meteorological variables provide auxiliary signals in some settings."
    )
    p = add_paragraph(doc, eng, first_line=False)
    for run in p.runs:
        set_run_font(run, "Times New Roman", 12)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    r = p.add_run("Key Words: ")
    set_run_font(r, "Times New Roman", 12, True)
    r = p.add_run("Influenza prediction; Multi-source data fusion; iTransformer; Time series; Baidu Index")
    set_run_font(r, "Times New Roman", 12)
    doc.add_page_break()

    add_toc(doc)

    # Section 1
    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_paragraph(doc, "流感是由流感病毒引起的急性呼吸道传染病，具有传播速度快、季节性强和易形成局部聚集性流行等特点。对流感活动趋势进行提前预测，有助于疾控部门和医疗机构提前识别高峰风险，合理安排发热门诊、检测试剂和抗病毒药物等资源。传统流感监测主要依赖哨点医院上报的流感样病例和实验室检测结果，这类监测数据可靠性较高，但往往存在一定滞后性。随着互联网搜索指数、气象再分析数据和公开卫生资料的积累，流感预测逐渐从单一监测序列建模转向多源异构数据融合建模。")
    add_paragraph(doc, "近年来，深度学习在时间序列预测领域取得了显著进展。循环神经网络和 LSTM 能够刻画时序依赖，但在长序列与多变量场景下存在训练效率低、变量间噪声相互干扰等问题。Transformer 及其变体通过注意力机制提升了全局依赖建模能力，其中 iTransformer 将每个变量的历史序列作为独立 token，在变量维度上建模跨变量关系，更适合气象、搜索指数和流感监测等异构变量的融合预测。")
    add_heading(doc, "1.2 研究目的与意义", 2)
    add_paragraph(doc, "本文的研究目标是构建一个可复现、可审计、可展示的周度流感活动趋势预测系统。系统以中国国家流感中心北方省份周度 ILI% 作为真实目标序列，融合北方代表城市气象数据和百度指数，完成未来 4 周流感活动水平预测。研究意义主要体现在三个方面：一是探索 iTransformer 在公共卫生多源时间序列预测中的适用性；二是形成从数据采集、质量审计到模型训练和 Web 展示的完整工程流程；三是用可追溯的官方周报数据替代旧版代理和虚拟延展目标序列，提高实验结论的数据可信度。")
    add_heading(doc, "1.3 国内外研究现状", 2)
    add_paragraph(doc, "传统流感预测方法主要包括 ARIMA、指数平滑、SIR 动力学模型和多元线性回归等。这些方法具有解释性强、计算成本低的优势，但通常依赖平稳性假设或简化传播机制，在面对非线性高峰、突发搜索热度变化和复杂气象影响时表现有限。机器学习方法如支持向量机、随机森林和梯度提升模型增强了非线性拟合能力，但仍需要较多人工特征工程。")
    add_paragraph(doc, "深度学习方法能够自动学习复杂时序模式。LSTM 等循环结构通过门控机制缓解长期依赖问题，但串行计算限制了训练效率。Transformer 结构通过自注意力机制建模全局关系，在交通、能源和金融时序预测中得到广泛应用。iTransformer 进一步将注意力建模的对象从时间点转向变量通道，降低了异构变量在同一时间步强行拼接带来的干扰。因此，将 iTransformer 用于流感病例数、气象要素和搜索指数的融合预测，具有较强的技术合理性。")
    add_heading(doc, "1.4 本文主要工作", 2)
    for item in [
        "构建三源数据处理流程：读取国家流感中心北方省份周报序列、北方代表城市 Open-Meteo 气象数据和百度指数，统一到周粒度。",
        "设计数据质量审计机制：记录数据来源、时间范围、缺失值、重复日期、异常相关性和解析状态。",
        "实现多步预测模型：以 iTransformer 为核心模型，并与 LSTM、DLinear、ARIMA 进行对比。",
        "开展分 horizon 评估和消融实验：分析未来第 1 至第 4 周误差变化，并比较不同数据源组合的贡献。",
        "实现 Web 可视化原型：展示数据概览、模型指标、预测曲线、注意力热力图和预警演示模块。",
    ]:
        add_paragraph(doc, item)

    # Section 2
    add_heading(doc, "2 数据来源与预处理", 1)
    add_heading(doc, "2.1 研究口径与数据边界", 2)
    add_paragraph(doc, "本文最终研究口径调整为中国北方省份周度流感活动预测。流感目标序列来自中国国家流感中心《中国流感监测周报》，以北方省份哨点医院报告的 ILI% 表征流感活动水平。该数据源长期公开发布，按周更新，较旧版平谷区代理序列具有更长时间跨度、更稳定口径和更强可追溯性。")
    add_paragraph(doc, f"国家流感中心原始周报序列覆盖 2010 年至 2026 年。考虑百度指数可用起点，本文多源建模使用 2011 年至 2026 年的重叠区间。当前周度合并数据共 {final_audit['merged_rows']} 周，经特征工程删除滞后与滚动窗口产生的前置缺失后，进入训练切分的数据共 {summary['split_metadata']['rows']} 周。个别周报 PDF 中实验室阳性率表格不可稳定抽取，因此 positive_rate 作为原始留痕字段保存，默认训练特征不使用该字段。")
    add_table(
        doc,
        "表2-1 数据来源与字段说明",
        ["数据类型", "文件路径", "时间范围", "主要字段", "说明"],
        [
            ["流感监测", "data/raw/flu/cnic_north_weekly_flu.csv", "2010-01-04至2026-04-13", "ili_rate, positive_rate", "国家流感中心周报；默认建模使用 ili_rate"],
            ["搜索指数", "data/raw/search/north_baidu_index.csv", "2011-01-01至2026-04-24", "flu/cold/fever_search_index", "北方代表城市百度指数日数据，按周聚合"],
            ["气象数据", "data/raw/weather/north_representative_city_weather.csv", "2011-01-01至2026-04-25", "temperature, humidity, wind_speed, pressure", "北方代表城市 Open-Meteo 日数据，按周聚合"],
        ],
        widths=[2.4, 4.8, 3.5, 4.5, 5.2],
    )
    add_picture(doc, FIG_DIR / "data_overview.png", "图2-1 多源数据时序概览", 6.2)
    add_heading(doc, "2.2 数据质量审计", 2)
    add_paragraph(doc, "为了避免数据来源不清和模拟数据误用，系统在训练前要求提供 source_manifest.json，记录每类数据的路径、来源名称、区域、粒度和采集方式。数据加载后，系统自动生成 data_quality_report.json，对原始数据和合并数据进行行数、字段、缺失值、重复日期、时间范围、数值统计和目标变量相关性检查。")
    add_paragraph(doc, f"最终复核显示，国家流感中心原始周报共 {final_audit['flu_rows']} 行，解析状态为 ok {final_audit['parse_status_counts'].get('ok', 0)} 周、partial {final_audit['parse_status_counts'].get('partial', 0)} 周；其中 ili_rate 原始缺失 {final_audit['missing_ili_rate_rows']} 周，positive_rate 原始缺失 {final_audit['missing_positive_rate_rows']} 周。缺失 ili_rate 的周次为 2017 年第 41 周、2017 年第 42 周和 2021 年第 50 周，预处理阶段采用线性插值与前后向填充处理目标序列边界和局部缺失。")
    add_paragraph(doc, "百度指数复核显示，flu_search_index、cold_search_index 和 fever_search_index 三个关键词每日均由 8 个北方代表地区参与聚合，region_count 全程 min=max=8。该结果说明当前 CSV 与北方代表城市聚合口径一致，但仍属于简单平均近似，不能解释为所有北方省份人口加权搜索行为。")
    merged = quality["datasets"]["merged"]
    add_table(
        doc,
        "表2-2 合并数据质量摘要",
        ["项目", "结果"],
        [
            ["样本数", str(merged["rows"])],
            ["时间范围", f"{merged['date_range']['start']} 至 {merged['date_range']['end']}"],
            ["重复日期", str(merged["duplicate_dates"])],
            ["是否按时间升序", "是" if merged["is_chronological"] else "否"],
            ["缺失值总数", str(sum(merged["missing_counts"].values()))],
            ["目标变量", "ili_rate"],
        ],
        widths=[4.0, 10.0],
    )
    add_heading(doc, "2.3 特征工程", 2)
    add_paragraph(doc, "在基础字段上，系统构造了时间周期特征、滞后特征、滚动统计特征、气象交互特征和搜索指数变化特征。时间特征采用 week_sin、week_cos、month_sin、month_cos 编码季节性；滞后特征包括 ili_rate 的 1、2、4 周滞后；滚动统计包括 4 周和 8 周均值与标准差；气象交互特征包括温湿度交互、舒适指数和风寒指数；搜索衍生特征包括关键词搜索指数环比变化率和加速度。")
    add_table(
        doc,
        "表2-3 特征类别说明",
        ["类别", "代表字段", "作用"],
        [
            ["流感历史", "ili_rate, lag, rolling", "刻画目标序列自身趋势与短期惯性"],
            ["时间周期", "week_sin, week_cos, month_sin, month_cos", "表达流感季节性周期"],
            ["气象特征", "temperature, humidity, wind_speed, pressure", "补充环境条件变化"],
            ["搜索指数", "flu_search_index, cold_search_index, fever_search_index", "反映公众症状关注与检索行为"],
        ],
        widths=[3.0, 6.5, 7.0],
    )
    add_picture(doc, FIG_DIR / "correlation_matrix.png", "图2-2 特征相关性矩阵", 5.8)

    # Section 3
    add_heading(doc, "3 模型与系统设计", 1)
    add_heading(doc, "3.1 总体技术路线", 2)
    add_paragraph(doc, "系统按照数据层、处理层、模型层、评估层和展示层组织。数据层负责保存原始 CSV 与来源清单；处理层完成清洗、周聚合、特征工程、归一化和时间切分；模型层包含 iTransformer 与基准模型；评估层输出整体指标、分 horizon 指标和图表；展示层基于 Streamlit 实现交互式展示。")
    add_picture(doc, FIG_DIR / "thesis_data_pipeline.png", "图3-1 多源数据处理与实验流程", 6.2)
    add_picture(doc, FIG_DIR / "thesis_system_architecture.png", "图3-2 系统总体架构", 6.2)
    add_heading(doc, "3.2 iTransformer 模型", 2)
    add_paragraph(doc, "iTransformer 的核心思想是将每个变量的历史窗口序列视为一个 token，在变量维度而非时间维度上计算自注意力。对于输入张量 X∈R^(B×C×L)，其中 B 表示批大小，C 表示变量数，L 表示回看窗口长度。模型首先通过通道独立嵌入层将每个变量的 L 维历史序列映射到 d_model 维空间，然后通过多层倒置 Transformer 编码器学习变量间关系，最后抽取目标变量 token 并映射为未来 4 周预测结果。")
    add_paragraph(doc, "本文设置 lookback_window 为 16 周，forecast_horizon 为 4 周，iTransformer 的 d_model 为 128，注意力头数为 8，编码层数为 2，前馈层维度为 256，dropout 为 0.1。训练采用 Adam 优化器、学习率 0.001、weight decay 0.01，并使用 early stopping 防止过拟合。")
    add_heading(doc, "3.3 基准模型", 2)
    add_paragraph(doc, "为验证 iTransformer 的预测效果，本文设置 LSTM、DLinear 和 ARIMA 作为基准模型。LSTM 用于代表循环神经网络类深度模型；DLinear 用于代表近年来长时序预测中常用的线性分解模型；ARIMA 用于代表传统单变量统计方法。所有深度模型使用相同的数据划分和目标变量，ARIMA 通过对齐滑动窗口方式输出与深度模型一致的多步预测点。")

    # Section 4
    add_heading(doc, "4 实验设计与结果分析", 1)
    add_heading(doc, "4.1 实验设置", 2)
    split = summary["split_metadata"]
    add_table(
        doc,
        "表4-1 数据集划分",
        ["集合", "样本行数", "时间范围"],
        [
            ["训练集", split["train_rows"], f"{split['train_date_range']['start']} 至 {split['train_date_range']['end']}"],
            ["验证集", split["val_rows"], f"{split['val_date_range']['start']} 至 {split['val_date_range']['end']}"],
            ["测试集", split["test_rows"], f"{split['test_date_range']['start']} 至 {split['test_date_range']['end']}"],
        ],
        widths=[3, 3, 7],
    )
    add_paragraph(doc, "评价指标包括 RMSE、MAE、MAPE、R2、峰值命中率、峰值时间偏移和峰值强度误差。其中 RMSE 和 MAE 衡量绝对误差，MAPE 衡量相对误差，R2 衡量解释方差比例，峰值相关指标用于观察模型对高发阶段的捕捉能力。")
    add_heading(doc, "4.2 多模型对比", 2)
    metrics = summary["metrics"]
    model_rows = []
    for name in ["iTransformer", "DLinear", "ARIMA", "LSTM"]:
        m = metrics[name]
        model_rows.append([name, fmt(m["RMSE"]), fmt(m["MAE"]), f"{fmt(m['MAPE'], 2)}%", fmt(m["R2"]), fmt(m["peak_hit_rate"], 3)])
    add_table(doc, "表4-2 多模型测试集指标对比", ["模型", "RMSE", "MAE", "MAPE", "R2", "峰值命中率"], model_rows, widths=[3, 2.2, 2.2, 2.2, 2.2, 2.6])
    add_paragraph(doc, f"结果显示，iTransformer 在当前真实北方周度数据版本下取得最优整体 RMSE，RMSE 为 {metrics['iTransformer']['RMSE']:.3f}，MAE 为 {metrics['iTransformer']['MAE']:.3f}，MAPE 为 {metrics['iTransformer']['MAPE']:.2f}%，R2 为 {metrics['iTransformer']['R2']:.3f}。相较 LSTM 和 DLinear，iTransformer 在整体误差上更低；ARIMA 在部分相对误差和峰值指标上仍具有竞争力，说明真实公共卫生序列中自回归结构较强。")
    add_picture(doc, FIG_DIR / "model_comparison.png", "图4-1 模型性能指标对比", 6.2)
    add_picture(doc, FIG_DIR / "multi_model_predictions.png", "图4-2 多模型预测趋势对比", 6.2)
    add_picture(doc, FIG_DIR / "iTransformer_predictions.png", "图4-3 iTransformer 预测结果与误差分布", 6.2)
    add_heading(doc, "4.3 多步预测分 horizon 分析", 2)
    h = summary["horizon_metrics"]["iTransformer"]
    add_table(
        doc,
        "表4-3 iTransformer 分 horizon 指标",
        ["预测步长", "RMSE", "MAE", "MAPE"],
        [
            ["第1周", fmt(h["H1_RMSE"]), fmt(h["H1_MAE"]), f"{fmt(h['H1_MAPE'], 2)}%"],
            ["第2周", fmt(h["H2_RMSE"]), fmt(h["H2_MAE"]), f"{fmt(h['H2_MAPE'], 2)}%"],
            ["第3周", fmt(h["H3_RMSE"]), fmt(h["H3_MAE"]), f"{fmt(h['H3_MAPE'], 2)}%"],
            ["第4周", fmt(h["H4_RMSE"]), fmt(h["H4_MAE"]), f"{fmt(h['H4_MAPE'], 2)}%"],
        ],
        widths=[3, 3, 3, 3],
    )
    add_paragraph(doc, "从分 horizon 结果看，未来第 1 周误差最低，随着预测距离增加，RMSE 和 MAE 呈上升趋势。这说明多步预测的不确定性会随预测步长累积，符合时间序列预测的一般规律。")
    add_heading(doc, "4.4 消融实验", 2)
    abl_sorted = sorted(ablation_rows, key=lambda x: float(x["RMSE"]))
    add_table(
        doc,
        "表4-4 iTransformer 多源特征消融实验",
        ["实验", "特征数", "RMSE", "MAE", "MAPE", "R2"],
        [[r["label"], r["feature_count"], fmt(r["RMSE"]), fmt(r["MAE"]), f"{fmt(r['MAPE'], 2)}%", fmt(r["R2"])] for r in abl_sorted],
        widths=[3.4, 2, 2.2, 2.2, 2.2, 2.2],
    )
    add_paragraph(doc, "消融实验固定模型结构为 iTransformer，仅改变输入特征组合。结果表明，在当前真实北方周度数据版本下，流感历史特征组合取得最低 RMSE，流感+搜索和三源融合结果较为接近。这说明换用真实目标序列后，旧版代理数据上的“搜索指数贡献最大”结论不能直接沿用；外生变量的贡献需要结合数据区间、流感季结构和特征聚合方式重新解释。")
    add_picture(doc, FIG_DIR / "ablation_comparison.png", "图4-4 iTransformer 消融实验指标对比", 6.2)
    add_heading(doc, "4.5 注意力可视化分析", 2)
    add_paragraph(doc, "本文保存了 iTransformer 最后一层的变量间注意力权重热力图，用于辅助观察模型在多源特征之间的关联建模情况。需要注意，注意力权重只能作为模型内部信息流的可视化参考，不能直接解释为流感传播的因果机制。")
    add_picture(doc, FIG_DIR / "attention_heatmap_layer0.png", "图4-5 iTransformer 变量间注意力权重热力图", 5.8)

    # Section 5
    add_heading(doc, "5 Web 展示系统实现", 1)
    add_heading(doc, "5.1 系统功能", 2)
    add_paragraph(doc, "项目实现了基于 Streamlit 的 Web 原型系统，主要页面包括数据总览、模型训练、预测分析、多模型对比实验、预警与调度演示和系统设置。数据总览页面展示多源时序、样本数量、时间范围和相关性矩阵；预测分析页面展示训练曲线、预测图和注意力热力图；对比实验页面展示多模型指标和消融实验结果；预警与调度页面以历史回测形式展示未来 4 周流感活动风险变化。")
    add_heading(doc, "5.2 工程复现性", 2)
    add_paragraph(doc, "系统采用配置文件 config/config.yaml 统一管理数据路径、模型参数、训练参数和报告路径。训练脚本 scripts/train.py 可一键完成数据采集、特征工程、预处理、训练、评估和报告导出；scripts/run_ablation.py 可复现实验消融。所有关键结果保存到 results/reports 与 results/figures，便于论文撰写和答辩展示。")

    # Section 6
    add_heading(doc, "6 不足与展望", 1)
    add_paragraph(doc, "第一，国家流感中心周报虽然公开可追溯，但个别早期或特殊年份的 PDF 表格不可稳定文本抽取，导致 positive_rate 字段存在缺失。本文将其作为原始留痕字段保存，正式训练主要依赖更稳定的 ili_rate 目标序列。")
    add_paragraph(doc, "第二，北方地区属于区域聚合口径，外生变量使用代表城市聚合近似，仍不能完全刻画所有北方省份的局部气象和搜索行为差异。后续可尝试按省份人口或哨点数量加权。")
    add_paragraph(doc, "第三，当前实验主要使用一次固定随机种子训练。后续可增加多随机种子实验、滚动时间切分和统计显著性检验，以增强结论稳定性。")
    add_paragraph(doc, "第四，Web 预警与调度模块目前属于原型展示，尚不能替代公共卫生业务系统。若用于真实场景，还需接入实时数据流、模型自动更新、告警阈值校准和人工审核机制。")

    # Section 7
    add_heading(doc, "7 结论", 1)
    add_paragraph(doc, "本文围绕周度流感活动趋势预测任务，完成了基于深度学习和多源数据融合的预测系统设计与实现。系统从数据来源清单、质量审计、特征工程、无泄漏时间切分、模型训练、评估报告到 Web 展示形成了完整工程闭环。实验结果表明，在国家流感中心北方省份真实周度监测数据上，iTransformer 相比 LSTM、DLinear 具有较优整体预测性能，并与 ARIMA 形成互补对照；消融实验显示流感历史特征仍是最稳定的信息来源。总体而言，本研究为流感趋势预测提供了一个结构完整、数据可信、可继续扩展的工程原型。")

    # References
    doc.add_page_break()
    add_center(doc, "参考文献", size=15, bold=True, font="黑体")
    refs = [
        "World Health Organization. Global Influenza Programme: Influenza surveillance and monitoring[EB/OL]. https://www.who.int/teams/global-influenza-programme/surveillance-and-monitoring.",
        "中国疾病预防控制中心病毒病预防控制所. 中国流感流行情况概要[EB/OL]. https://ivdc.chinacdc.cn/cnic/zyzx/lgzb/.",
        "中国国家流感中心. 中国流感监测周报[EB/OL]. https://ivdc.chinacdc.cn/cnic/zyzx/lgzb/.",
        "Liu Y, Hu T, Zhang H, et al. iTransformer: Inverted Transformers Are Effective for Time Series Forecasting[C]. International Conference on Learning Representations, 2024.",
        "Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems, 2017:5998-6008.",
        "Hochreiter S, Schmidhuber J. Long Short-Term Memory[J]. Neural Computation, 1997,9(8):1735-1780.",
        "Box G E P, Jenkins G M, Reinsel G C, Ljung G M. Time Series Analysis: Forecasting and Control[M]. 5th ed. Hoboken: Wiley, 2015.",
        "Hyndman R J, Athanasopoulos G. Forecasting: Principles and Practice[M/OL]. 3rd ed. OTexts, 2021.",
        "Zeng A, Chen M, Zhang L, Xu Q. Are Transformers Effective for Time Series Forecasting?[C]. AAAI Conference on Artificial Intelligence, 2023.",
        "Open-Meteo. Historical Weather API Documentation[EB/OL]. https://open-meteo.com/en/docs/historical-weather-api.",
        "百度指数. 百度指数平台[EB/OL]. https://index.baidu.com/.",
        "McKinney W. Data Structures for Statistical Computing in Python[C]. Proceedings of the 9th Python in Science Conference, 2010:56-61.",
        "Paszke A, Gross S, Massa F, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library[C]. Advances in Neural Information Processing Systems, 2019:8024-8035.",
        "Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine Learning in Python[J]. Journal of Machine Learning Research, 2011,12:2825-2830.",
        "Seabold S, Perktold J. Statsmodels: Econometric and Statistical Modeling with Python[C]. Proceedings of the 9th Python in Science Conference, 2010:92-96.",
        "Plotly Technologies Inc. Collaborative data science with Plotly[EB/OL]. https://plotly.com/.",
    ]
    for i, ref in enumerate(refs, start=1):
        add_paragraph(doc, f"[{i}] {ref}", first_line=False)

    # Acknowledgements
    doc.add_page_break()
    add_center(doc, "致  谢", size=15, bold=True, font="黑体")
    add_paragraph(doc, "在毕业设计完成过程中，指导教师在选题论证、研究路线、实验设计和论文撰写等方面给予了耐心指导，使我能够逐步明确项目边界并完成系统实现。感谢信息科学与工程学部提供的学习环境和课程训练，使我具备了开展数据处理、深度学习建模和系统开发的基础能力。")
    add_paragraph(doc, "同时感谢开源社区提供的 PyTorch、Pandas、Statsmodels、Streamlit 等工具，这些工具为本文系统开发和实验复现提供了重要支持。由于个人能力和数据条件有限，本文仍存在区域外生变量聚合较粗、部分 PDF 字段抽取不完整等问题，后续将继续完善数据来源与模型验证工作。")

    # Appendix
    doc.add_page_break()
    add_center(doc, "附  录", size=15, bold=True, font="黑体")
    add_heading(doc, "附录A 主要运行命令", 1)
    add_paragraph(doc, "完整训练命令：C:\\ProgramData\\anaconda3\\envs\\ocean_torch\\python.exe scripts\\train.py", first_line=False)
    add_paragraph(doc, "消融实验命令：C:\\ProgramData\\anaconda3\\envs\\ocean_torch\\python.exe scripts\\run_ablation.py --skip-collect", first_line=False)
    add_paragraph(doc, "Web 启动命令：C:\\ProgramData\\anaconda3\\envs\\ocean_torch\\python.exe -m streamlit run web\\app.py", first_line=False)
    add_heading(doc, "附录B 主要输出文件", 1)
    for item in [
        "results/reports/experiment_summary.json",
        "results/reports/horizon_metrics.json",
        "results/reports/ablation_report.md",
        "results/reports/data_quality_report.json",
        "results/reports/final_data_audit.md",
        "results/figures/model_comparison.png",
        "results/figures/ablation_comparison.png",
    ]:
        add_paragraph(doc, item, first_line=False)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
