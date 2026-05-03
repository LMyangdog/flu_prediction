from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT = PROJECT_ROOT / "毕设答辩知识准备手册.docx"
FALLBACK_OUT = PROJECT_ROOT / "毕设答辩知识准备手册_项目建设过程版.docx"

ACCENT = "1F4E79"
ACCENT_DARK = "17365D"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F7"
LIGHT_GREEN = "EAF7EA"
LIGHT_YELLOW = "FFF7DF"
TEXT = "222222"
MUTED = "666666"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)


def set_run_font(run, east_asia: str = "微软雅黑", ascii_font: str = "Aptos") -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_font(paragraph, size: float = 10.5, color: str = TEXT) -> None:
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color, before, after in [
        ("Title", 24, ACCENT_DARK, 0, 12),
        ("Heading 1", 17, ACCENT_DARK, 16, 8),
        ("Heading 2", 13.5, ACCENT, 10, 5),
        ("Heading 3", 11.5, TEXT, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Heading 3"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name in {"Heading 1", "Heading 2"}


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(78)
    r = p.add_run("基于深度学习和多元数据的流感爆发趋势预测")
    set_run_font(r)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("毕设答辩知识准备手册")
    set_run_font(r)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("从基础概念到答辩问答的复习版")
    set_run_font(r)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    info = [
        ("研究对象", "中国北方省份周度流感活动趋势"),
        ("预测目标", "国家流感中心周报中的 ILI%（ili_rate）"),
        ("核心模型", "iTransformer 多变量时间序列预测模型"),
        ("复习用途", "用于答辩前知识梳理、后续提问与持续修改"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(info):
        row = table.rows[i]
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(11.5)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_shading(row.cells[1], "FFFFFF")
        set_cell_text(row.cells[0], k, bold=True, color=ACCENT_DARK)
        set_cell_text(row.cells[1], v)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("建议使用方式：先通读本手册，再按“概念不清楚 -> 模型讲不出 -> 指标不会解释 -> 问答不熟练”的顺序逐段复习。")
    set_run_font(r)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_YELLOW) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r)
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"- {item}")
        set_run_font(r)
        r.font.size = Pt(10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{idx}. {item}")
        set_run_font(r)
        r.font.size = Pt(10.5)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        if widths:
            hdr[i].width = Cm(widths[i])
        set_cell_shading(hdr[i], ACCENT)
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            if widths:
                cells[i].width = Cm(widths[i])
            set_cell_shading(cells[i], "FFFFFF" if len(table.rows) % 2 else LIGHT_GRAY)
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line_no, line in enumerate(text.splitlines()):
        if line_no:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph()


def add_project_overview(doc: Document) -> None:
    doc.add_heading("1. 项目总览：先把故事讲清楚", level=1)
    add_callout(
        doc,
        "一句话版本",
        "本项目利用过去 16 周的流感监测、气象和百度搜索指数等 29 个周度特征，预测未来 1-4 周中国北方地区的 ILI% 流感活动趋势，并用 iTransformer 与 LSTM、DLinear、ARIMA 进行对比。",
        LIGHT_GREEN,
    )
    doc.add_heading("1.1 预测对象是什么", level=2)
    p = doc.add_paragraph()
    p.add_run("预测目标是 ").bold = False
    r = p.add_run("ili_rate，也就是 ILI%")
    r.bold = True
    set_run_font(r)
    p.add_run("。ILI 是 Influenza-like Illness，中文通常称为“流感样病例”。ILI% 可以理解为流感样病例在门急诊就诊人群中的占比，用来反映流感活动水平。")
    set_paragraph_font(p)
    add_code_block(doc, "ILI% = 流感样病例数 / 门急诊总就诊人数 × 100%")
    add_bullets(doc, [
        "它不是确诊流感人数，而是公共卫生监测中常用的流感活动强度指标。",
        "本文研究口径是中国北方省份周度 ILI%，不能外推成单一城市或区县的确诊病例预测。",
        "positive_rate 是流感病毒检测阳性率，本项目中主要用于原始监测留痕和数据质量说明，默认不进入正式训练特征。",
    ])

    doc.add_heading("1.2 数据来源与研究口径", level=2)
    add_simple_table(
        doc,
        ["数据类型", "项目中的作用", "答辩表述要点"],
        [
            ["流感监测数据", "提供目标变量 ili_rate", "来自国家流感中心《中国流感监测周报》，公开可追溯。"],
            ["气象数据", "提供温度、湿度、风速、气压等外生变量", "气象因素会影响流感传播环境，按周聚合后与 ILI% 对齐。"],
            ["百度搜索指数", "提供公众搜索行为和症状关注度变量", "搜索热度可能在一定程度上反映公众症状或关注变化。"],
        ],
        [3.2, 5.5, 7.2],
    )
    add_callout(
        doc,
        "答辩安全表述",
        "本文预测的是国家流感中心北方省份周度 ILI%，不是某个城市的真实确诊病例数；研究结论应限定在当前数据口径、时间切分和特征聚合方式下。",
        LIGHT_YELLOW,
    )


def add_time_series_basics(doc: Document) -> None:
    doc.add_heading("2. 时间序列预测基础", level=1)
    doc.add_heading("2.1 什么是时间序列", level=2)
    p = doc.add_paragraph()
    p.add_run("时间序列是按时间顺序排列的数据。你的项目中，ILI% 每周都有一个观测值，因此构成周度时间序列。模型学习的是历史变化规律，并尝试预测未来趋势。")
    set_paragraph_font(p)
    add_code_block(doc, "第1周 ILI% = 2.1\n第2周 ILI% = 2.3\n第3周 ILI% = 2.8\n...\n输入过去16周 → 预测未来4周")
    doc.add_heading("2.2 滑动窗口如何把时间序列变成训练样本", level=2)
    p = doc.add_paragraph()
    p.add_run("深度学习模型通常需要输入 X 和标签 Y。时间序列预测中常用滑动窗口方法：用连续 16 周作为输入，再取其后的 4 周作为预测标签。窗口不断向后滑动，就能构造出多个训练样本。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["概念", "项目取值", "含义"],
        [
            ["lookback_window", "16", "模型每次看过去 16 周历史信息。"],
            ["forecast_horizon", "4", "模型一次预测未来 4 周 ILI%。"],
            ["H1/H2/H3/H4", "未来第1-4周", "用于分析不同预测步长上的误差。"],
            ["多步预测", "一次输出4个值", "既能看短期预警，也能看较远趋势。"],
        ],
        [4.0, 3.0, 8.5],
    )
    add_callout(
        doc,
        "可以直接背",
        "本研究采用滑动窗口方法，将连续周度数据转换为监督学习样本。模型输入为过去 16 周的多变量序列，输出为未来 4 周的 ILI% 预测值。",
        LIGHT_GREEN,
    )


def add_feature_engineering(doc: Document) -> None:
    doc.add_heading("3. 多源数据融合与特征工程", level=1)
    doc.add_heading("3.1 为什么要融合多源数据", level=2)
    p = doc.add_paragraph()
    p.add_run("单一流感历史序列只能反映自身变化，而气象和搜索指数可以提供外部环境与公众行为信息。流感活动通常具有季节性，并受到温湿度、人群活动和公众关注度变化影响，因此多源数据融合有助于模型获得更丰富的上下文。")
    set_paragraph_font(p)
    add_callout(
        doc,
        "注意措辞",
        "不要说“气象和搜索指数一定提升效果”。更稳妥的说法是：外生变量的作用受到聚合口径、噪声和流行季阶段影响，本文通过消融实验进行了对比分析。",
        LIGHT_YELLOW,
    )

    doc.add_heading("3.2 项目中的 29 个特征如何理解", level=2)
    add_simple_table(
        doc,
        ["类别", "代表特征", "作用"],
        [
            ["原始监测/外生特征", "ili_rate、temperature、humidity、wind_speed、pressure、搜索指数", "提供目标历史、环境信息和公众关注度信息。"],
            ["周期特征", "week_sin、week_cos、month_sin、month_cos、is_flu_season", "表达流感活动的季节性和周期性。"],
            ["滞后特征", "ili_rate_lag1、lag2、lag4", "表达前 1、2、4 周流感水平对当前和未来的影响。"],
            ["滚动统计", "4/8 周滚动均值、滚动标准差", "描述近期趋势水平和波动程度。"],
            ["交互与变化特征", "温湿度交互、舒适指数、搜索变化率、搜索加速度", "补充气象组合效应和搜索热度变化速度。"],
        ],
        [3.6, 5.6, 6.5],
    )
    doc.add_heading("3.3 为什么周期特征要用 sin/cos", level=2)
    p = doc.add_paragraph()
    p.add_run("周数和月份具有循环性。第 52 周和第 1 周在时间上很接近，如果直接用数字表示，模型会误以为 52 和 1 差距很大。sin/cos 编码可以把周期首尾自然连接起来，更适合表达季节规律。")
    set_paragraph_font(p)


def add_project_build_process(doc: Document) -> None:
    doc.add_heading("4. 项目建设过程：从架构搭建到最终实现", level=1)
    add_callout(
        doc,
        "这一节的答辩作用",
        "当老师问“你这个系统是怎么做出来的”时，不要只说用了 iTransformer，而要按工程链路回答：先搭目录和配置，再采集真实数据，做周度对齐与特征工程，最后训练模型、评估结果并接入 Web 展示。",
        LIGHT_GREEN,
    )

    doc.add_heading("4.1 项目架构如何设计", level=2)
    p = doc.add_paragraph()
    p.add_run("项目采用“配置、数据、模型、训练、结果、展示”分层组织。这样做的好处是每一层职责清晰：配置文件控制实验参数，数据模块负责采集和处理，模型模块只关心网络结构，训练模块统一训练与评估，Web 模块读取结果用于展示。")
    set_paragraph_font(p)
    add_code_block(
        doc,
        "flu_prediction/\n"
        "├── config/config.yaml          # 全局配置：数据路径、特征列、模型和训练参数\n"
        "├── data/raw                    # 原始流感、气象、搜索指数数据\n"
        "├── data/processed              # 周度融合数据与特征工程后数据\n"
        "├── data/splits                 # 训练/验证/测试切分后的 numpy 数据\n"
        "├── scripts                     # 采集、训练、消融、审计等运行脚本\n"
        "├── src/data                    # 数据采集、预处理、特征工程、质量审计\n"
        "├── src/models                  # iTransformer、LSTM、DLinear、ARIMA\n"
        "├── src/training                # 训练器、早停、调度、检查点保存\n"
        "├── src/utils                   # 指标计算与可视化\n"
        "├── results                     # 图表、日志、实验报告\n"
        "└── web/app.py                  # Streamlit 可视化展示系统",
    )
    add_simple_table(
        doc,
        ["层次", "关键文件/目录", "主要作用"],
        [
            ["配置层", "config/config.yaml", "集中管理数据路径、特征、窗口长度、模型参数和训练参数。"],
            ["数据层", "src/data、data/", "完成真实数据读取、周度对齐、质量审计、缺失处理和切分。"],
            ["模型层", "src/models", "封装 iTransformer 及对比模型，保证实验可复现。"],
            ["训练评估层", "scripts/train.py、src/training", "统一执行训练、验证、测试、指标计算和结果保存。"],
            ["展示层", "web/app.py、results/", "读取实验产物，展示数据概览、预测曲线、指标和结论。"],
        ],
        [3.0, 4.7, 8.2],
    )

    doc.add_heading("4.2 数据采集如何构建", level=2)
    p = doc.add_paragraph()
    p.add_run("数据采集阶段的目标不是简单下载文件，而是建立一个可追溯的数据入口。项目通过 source_manifest.json 登记数据来源、区域、粒度和文件路径，并启用严格真实数据模式，避免实验自动回退到模拟数据。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["数据源", "采集/整理方式", "进入项目后的处理"],
        [
            ["流感监测", "scripts/fetch_flu_cn_weekly.py 解析国家流感中心周报 HTML/PDF。", "整理为 cnic_north_weekly_flu.csv，提供 ili_rate。"],
            ["气象数据", "Open-Meteo Archive API 获取 8 个北方代表城市日度数据。", "按日期取均值，再按 ISO 周聚合。"],
            ["搜索指数", "百度指数登录导出或脚本整理为标准 CSV。", "保留“流感、感冒、发烧”等关键词日度指数，再聚合到周。"],
        ],
        [3.0, 6.2, 6.7],
    )
    add_bullets(doc, [
        "MultiSourceDataCollector 统一读取三类数据，并检查必要字段是否存在。",
        "日度气象和搜索指数先转换为 year/week，再按周求均值。",
        "三类数据最终按 year/week 合并，输出 data/processed/merged_dataset.csv。",
        "同时输出 data_quality_report.json，用于论文和答辩说明数据可信性。",
    ])

    doc.add_heading("4.3 数据处理与防止信息泄漏", level=2)
    p = doc.add_paragraph()
    p.add_run("时间序列预测最重要的工程细节之一是防止未来信息泄漏。项目在预处理阶段按时间顺序切分训练集、验证集和测试集，并且归一化只在训练集上拟合，再应用到验证集和测试集。这样测试集不会提前参与模型训练或数据标准化。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["处理步骤", "项目做法", "为什么重要"],
        [
            ["日期排序", "统一 date 格式并按时间升序排列。", "保证模型只利用历史信息。"],
            ["缺失值处理", "对训练特征线性插值，再前向/后向填充边界。", "避免少量缺失导致样本无法训练。"],
            ["异常值处理", "使用 IQR 方法裁剪极端值。", "降低异常尖点对模型参数的扰动。"],
            ["严格切分", "按时间顺序划分 train/val/test。", "避免随机打乱破坏时间因果关系。"],
            ["归一化", "MinMaxScaler 只在训练集 fit。", "防止验证集和测试集分布泄漏。"],
            ["滑动窗口", "每个集合内独立切片，生成 X 和 y。", "避免窗口跨集合边界造成信息重叠。"],
        ],
        [3.0, 6.3, 6.6],
    )
    add_code_block(
        doc,
        "预处理输出形状：\n"
        "X: (num_samples, num_variables, lookback_window)\n"
        "y: (num_samples, forecast_horizon)\n\n"
        "在本项目中可理解为：\n"
        "X: (样本数, 29, 16)\n"
        "y: (样本数, 4)",
    )

    doc.add_heading("4.4 特征工程如何接入训练", level=2)
    p = doc.add_paragraph()
    p.add_run("特征工程由 src/data/feature_engineer.py 统一完成，核心思路是把“原始观测值”扩展成更利于模型学习的时序特征。它不是脱离模型的附加工作，而是直接决定模型输入通道数量和信息质量。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["特征工程环节", "新增信息", "服务的建模问题"],
        [
            ["时间特征", "week/month 的 sin/cos、流感季标志", "让模型感知季节周期。"],
            ["滞后特征", "前 1、2、4 周 ILI%", "利用流感活动的自相关性。"],
            ["滚动统计", "4/8 周均值和标准差", "表达近期趋势水平和波动。"],
            ["气象交互", "温湿度交互、舒适指数、风寒指数", "表达气象变量的组合影响。"],
            ["搜索衍生", "搜索指数变化率和加速度", "捕捉公众关注度变化速度。"],
        ],
        [3.2, 5.5, 7.0],
    )

    doc.add_heading("4.5 模型训练与结果产出", level=2)
    p = doc.add_paragraph()
    p.add_run("训练入口是 scripts/train.py。脚本会读取配置、设置随机种子、准备数据、创建 DataLoader，然后分别训练 iTransformer、LSTM、DLinear，并运行 ARIMA 基线。训练器统一使用 MSELoss、Adam 优化器、余弦学习率调度、梯度裁剪和早停机制。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["阶段", "关键实现", "产物"],
        [
            ["数据加载", "FluDataset + DataLoader", "训练/验证/测试批数据。"],
            ["模型构建", "build_* 模型构建函数", "可训练模型实例。"],
            ["训练控制", "Trainer + EarlyStopping", "最佳模型权重和训练历史。"],
            ["测试评估", "metrics 模块", "RMSE、MAE、MAPE、R²、H1-H4 指标。"],
            ["可视化报告", "Visualizer + JSON 报告", "预测曲线、对比图、实验简报和答辩材料依据。"],
        ],
        [3.0, 6.0, 6.7],
    )
    add_code_block(
        doc,
        "常用运行命令：\n"
        "python scripts/train.py\n"
        "python scripts/train.py --skip-collect\n"
        "python scripts/run_ablation.py --skip-collect\n"
        "streamlit run web/app.py",
    )

    doc.add_heading("4.6 最终实现与展示系统", level=2)
    p = doc.add_paragraph()
    p.add_run("最终系统不是只停留在模型训练，而是把实验结果落到可查看的 Web 页面和答辩材料中。web/app.py 使用 Streamlit 读取 results/reports 与 results/figures 中的产物，展示数据概览、模型指标、预测曲线、预警复盘和附录信息。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["展示内容", "读取产物", "答辩用途"],
        [
            ["数据概览", "merged_dataset + audit", "说明数据来源、规模和质量。"],
            ["模型指标", "summary + horizon_metrics", "展示模型对比和多步误差。"],
            ["预测曲线", "test_predictions + figures", "直观看真实值与预测值走势。"],
            ["消融实验", "ablation summaries", "解释不同数据源的贡献。"],
            ["附录说明", "reports/ 下的审计与简报", "支撑老师追问时的数据可信性说明。"],
        ],
        [3.2, 5.6, 6.9],
    )
    add_callout(
        doc,
        "项目建设过程答辩话术",
        "本项目先按功能分层搭建工程目录，再建立真实数据采集和来源登记机制；之后将流感、气象、搜索指数统一到周粒度，完成缺失处理、异常值处理、特征工程和严格时间切分；最后用统一训练器完成多模型训练、指标评估、图表导出，并通过 Streamlit 页面展示预测结果。",
        LIGHT_GREEN,
    )


def add_itransformer(doc: Document) -> None:
    doc.add_heading("5. iTransformer：答辩必须讲清楚的核心模型", level=1)
    doc.add_heading("5.1 从 Transformer 到 iTransformer", level=2)
    p = doc.add_paragraph()
    p.add_run("传统 Transformer 最早用于自然语言处理，通常把句子中的每个词看作一个 token。用于时间序列时，常见做法是把每个时间点作为 token，每个 token 中包含多个变量。iTransformer 的思路是倒置：不把时间点作为 token，而是把每个变量的完整历史序列作为 token。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["建模方式", "Token 是什么", "更关注什么", "适用说明"],
        [
            ["普通时间序列 Transformer", "每个时间点", "时间点之间的关系", "适合强调长时间依赖的场景。"],
            ["iTransformer", "每个变量", "变量之间的关系", "适合多变量、多源异构时间序列。"],
        ],
        [4.2, 4.2, 4.2, 4.2],
    )
    add_callout(
        doc,
        "核心理解",
        "iTransformer 的“倒置”不是把时间倒过来，而是把建模视角从“时间点之间的注意力”转为“变量之间的注意力”。",
        LIGHT_GREEN,
    )

    doc.add_heading("5.2 你的项目中 iTransformer 的输入输出", level=2)
    add_code_block(doc, "输入形状：(B, 29, 16)\nB：batch size\n29：变量数\n16：过去16周历史\n\nChannel Embedding 后：(B, 29, 128)\n倒置注意力编码后：(B, 29, 128)\n取目标变量 ili_rate 的 token\n输出形状：(B, 4)")
    add_bullets(doc, [
        "每个变量的 16 周历史先被映射成一个 128 维变量 token。",
        "29 个变量 token 进入倒置多头自注意力层，模型学习变量之间的相关性。",
        "最终取目标变量 ili_rate 对应的 token，投影得到未来 4 周预测值。",
    ])

    doc.add_heading("5.3 通道独立嵌入和倒置注意力", level=2)
    add_simple_table(
        doc,
        ["模块", "通俗解释", "在项目中的意义"],
        [
            ["通道独立嵌入", "每个变量先单独处理自己的历史序列。", "避免温度、湿度、搜索指数、ILI% 在同一时间点被粗糙拼接造成噪声干扰。"],
            ["倒置多头注意力", "让变量之间互相“看见”并学习影响权重。", "模型可学习 ILI%、季节、气象和搜索指数之间的复杂关系。"],
            ["目标 token 投影", "只用 ili_rate 的综合表示输出预测。", "减少把所有变量直接展开带来的过拟合风险。"],
        ],
        [3.6, 5.8, 6.4],
    )
    add_callout(
        doc,
        "答辩版表述",
        "iTransformer 将每个变量的完整历史序列作为 token，通过倒置注意力学习变量间关系，更适合本文这种流感、气象、搜索指数融合的多源时间序列预测任务。",
        LIGHT_GREEN,
    )

    doc.add_heading("5.4 注意力机制怎么讲", level=2)
    p = doc.add_paragraph()
    p.add_run("注意力机制可以理解为：模型在预测时自动判断哪些变量更重要。例如预测流感趋势时，模型可能关注近期 ILI% 的变化、是否处于流感季、搜索“发烧”的热度变化，以及温湿度等气象条件。多头注意力则相当于从多个角度观察变量关系。")
    set_paragraph_font(p)
    add_callout(
        doc,
        "一句话解释多头注意力",
        "多头注意力机制允许模型从多个子空间学习变量之间的不同关联模式，从而增强对复杂非线性关系的表达能力。",
        LIGHT_BLUE,
    )


def add_baselines_metrics(doc: Document) -> None:
    doc.add_heading("6. 对比模型与评价指标", level=1)
    doc.add_heading("6.1 为什么要设置对比模型", level=2)
    p = doc.add_paragraph()
    p.add_run("只报告一个模型的结果，很难说明它到底好不好。因此需要用传统统计模型、经典深度学习模型和强基线模型进行对比，证明核心模型在当前任务上的相对效果。")
    set_paragraph_font(p)
    add_simple_table(
        doc,
        ["模型", "类型", "优点", "局限"],
        [
            ["ARIMA", "传统统计模型", "理论成熟，适合规律性较强的单变量序列。", "不擅长多源非线性特征融合。"],
            ["LSTM", "循环神经网络", "能学习时间依赖，曾是经典序列模型。", "串行计算较慢，长序列可能存在梯度衰减和训练不稳定。"],
            ["DLinear", "线性时序预测基线", "结构简单，对趋势和季节性有较强基线意义。", "复杂非线性关系表达能力有限。"],
            ["iTransformer", "多变量深度时序模型", "适合学习变量间关系，适配多源融合。", "需要训练数据和超参数支持，解释性弱于传统统计模型。"],
        ],
        [2.7, 3.0, 5.7, 5.7],
    )

    doc.add_heading("6.2 RMSE、MAE、MAPE、R²", level=2)
    add_simple_table(
        doc,
        ["指标", "含义", "怎么看", "答辩解释"],
        [
            ["RMSE", "均方根误差", "越低越好，对大误差更敏感。", "适合观察高峰期预测偏差，因为大误差会被放大。"],
            ["MAE", "平均绝对误差", "越低越好，直观稳定。", "表示平均相差多少个 ILI 百分点。"],
            ["MAPE", "平均绝对百分比误差", "越低越好，用百分比表示。", "便于非技术听众理解平均相对误差。"],
            ["R²", "决定系数", "越接近 1 越好；小于 0 表示不如均值基线。", "衡量模型对真实波动的解释能力。"],
        ],
        [2.4, 4.2, 4.3, 5.7],
    )
    add_code_block(doc, "RMSE = sqrt(mean((真实值 - 预测值)^2))\nMAE  = mean(abs(真实值 - 预测值))\nMAPE = mean(abs((真实值 - 预测值) / 真实值)) × 100%\nR²   = 1 - 残差平方和 / 总平方和")
    add_callout(
        doc,
        "指标背诵版",
        "RMSE、MAE、MAPE 越低越好，R² 越高越好。RMSE 更惩罚大误差，MAE 更直观，MAPE 适合解释相对误差，R² 表示模型对波动的解释程度。",
        LIGHT_GREEN,
    )


def add_results(doc: Document) -> None:
    doc.add_heading("7. 项目实验结果怎么讲", level=1)
    doc.add_heading("7.1 总体模型对比", level=2)
    add_simple_table(
        doc,
        ["模型", "RMSE", "MAE", "MAPE", "R²"],
        [
            ["iTransformer", "0.772", "0.491", "11.51%", "0.538"],
            ["ARIMA", "0.980", "0.552", "11.87%", "0.450"],
            ["DLinear", "1.031", "0.694", "15.73%", "0.176"],
            ["LSTM", "1.366", "0.880", "18.81%", "-0.444"],
        ],
        [4.2, 2.8, 2.8, 3.0, 2.8],
    )
    add_bullets(doc, [
        "iTransformer 的整体 RMSE 和 MAE 最低，说明当前测试集上的总体误差较小。",
        "ARIMA 作为传统统计模型仍有竞争力，说明周度 ILI% 序列存在较强季节性和自相关性。",
        "DLinear 能捕捉部分线性趋势，但对多源非线性关系的表达能力有限。",
        "LSTM 在当前设置下 R² 为负，表示其测试集表现不如简单均值基线。",
    ])
    add_callout(
        doc,
        "稳妥结论",
        "在当前北方地区真实周度 ILI% 测试集上，iTransformer 取得最低的整体 RMSE 和 MAE，说明其整体拟合误差较小；ARIMA 在部分短期预测和峰值相关指标上仍具有一定竞争力。",
        LIGHT_YELLOW,
    )

    doc.add_heading("7.2 多步预测结果", level=2)
    add_simple_table(
        doc,
        ["预测步长", "RMSE", "MAE", "MAPE"],
        [
            ["H1", "0.525", "0.351", "8.38%"],
            ["H2", "0.676", "0.459", "11.27%"],
            ["H3", "0.828", "0.535", "12.48%"],
            ["H4", "0.983", "0.619", "13.89%"],
        ],
        [4.0, 3.5, 3.5, 3.5],
    )
    p = doc.add_paragraph()
    p.add_run("解释重点：H1 到 H4 的误差逐渐增大，这是多步时间序列预测中的常见现象。预测越远，未知因素越多，模型不确定性越高。")
    set_paragraph_font(p)

    doc.add_heading("7.3 消融实验怎么解释", level=2)
    add_simple_table(
        doc,
        ["实验组", "特征数量", "作用"],
        [
            ["仅流感历史", "13", "检验历史 ILI%、季节项、滞后和滚动特征的基础预测能力。"],
            ["流感 + 气象", "20", "检验温度、湿度、风速、气压及气象交互项的贡献。"],
            ["流感 + 搜索", "22", "检验百度搜索指数及其变化率、加速度的贡献。"],
            ["三源融合", "29", "综合利用流感、气象和搜索指数特征。"],
        ],
        [3.2, 2.4, 10.6],
    )
    add_callout(
        doc,
        "消融实验答辩话术",
        "消融实验用于分析不同数据源对预测性能的贡献，而不是简单证明某一类外生变量必然有效。外生变量可能带来辅助信息，也可能因为聚合噪声、时间滞后或流行季差异而收益有限。",
        LIGHT_BLUE,
    )


def add_qa(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("8. 答辩高频问答", level=1)
    qas = [
        ("为什么用周度数据，不用日度数据？", "国家流感中心公开监测数据主要以周为单位发布。若强行构造日度 ILI%，会引入插值或伪目标，降低实验可信度。因此本文将日度气象和搜索指数聚合到周度，与真实流感监测口径对齐。"),
        ("为什么用北方地区，不用单个城市？", "国家流感中心长期公开南北方省份周度 ILI% 数据，时间跨度长、口径稳定、可追溯性强。相比单一区县或城市代理数据，北方地区序列更适合作为毕业设计的真实实验目标。"),
        ("为什么 iTransformer 适合这个项目？", "本项目输入变量来自流感、气象和搜索指数，具有多源异构特征。iTransformer 将每个变量的历史序列作为 token，通过倒置注意力学习变量间关系，能更自然地建模多变量之间的关联。"),
        ("模型是不是能预测真实确诊人数？", "不能这样说。本文预测的是 ILI%，也就是流感样病例占比，用于反映流感活动趋势，不等同于确诊病例数。"),
        ("为什么要做消融实验？", "消融实验用于分析不同数据源对预测性能的贡献。例如只用流感历史、加入气象、加入搜索指数、三源融合分别训练模型，从而判断外生变量是否带来有效信息。"),
        ("为什么 LSTM 表现不如 iTransformer？", "LSTM 按时间步串行建模，面对多源异构变量时容易受到特征噪声和训练稳定性的影响。iTransformer 从变量维度学习关系，更契合本项目的多变量融合任务。"),
        ("为什么 ARIMA 也表现不错？", "ILI% 是周度公共卫生监测序列，具有一定季节性和自相关性，因此传统统计模型在短期预测上仍有竞争力。这并不否定 iTransformer，而是说明基线选择具有参考价值。"),
    ]
    for q, a in qas:
        h = doc.add_heading(q, level=2)
        h.paragraph_format.keep_with_next = True
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        r = p.add_run("答：")
        r.bold = True
        set_run_font(r)
        p.add_run(a)
        set_paragraph_font(p)


def add_memory_checklist(doc: Document) -> None:
    doc.add_heading("9. 答辩前必须背熟的句子", level=1)
    add_numbered(doc, [
        "本文预测目标是国家流感中心北方省份周度 ILI%，不是确诊病例数。",
        "输入为过去 16 周的 29 个多源特征，输出为未来 4 周 ILI%。",
        "iTransformer 的核心是倒置建模，把变量作为 token，用注意力学习变量间关系。",
        "RMSE、MAE、MAPE 越低越好，R² 越高越好。",
        "H1 到 H4 误差逐渐增大，说明预测步长越远，不确定性越高。",
        "消融实验用于分析数据源贡献，不能简单说气象或搜索指数一定有效。",
        "项目结论限定在当前北方地区真实周度监测数据、特征聚合和时间切分口径下。",
    ])
    doc.add_heading("9.1 后续可继续补充的问题", level=2)
    add_bullets(doc, [
        "iTransformer 原论文与普通 Transformer 的更细区别。",
        "注意力机制、残差连接、LayerNorm、Dropout、Adam 优化器等深度学习基础。",
        "训练集、验证集、测试集为什么要按时间顺序划分。",
        "数据补齐、缺失值处理、搜索指数聚合口径如何答辩。",
        "Web 系统展示部分如何从工程实现角度介绍。",
    ])


def add_references(doc: Document) -> None:
    doc.add_heading("参考与项目依据", level=1)
    add_bullets(doc, [
        "项目 README、实验报告与配置文件：D:\\flu_prediction\\README.md、results\\reports、config\\config.yaml。",
        "iTransformer 原论文：Liu et al., iTransformer: Inverted Transformers Are Effective for Time Series Forecasting, ICLR 2024。",
        "项目论文与答辩材料口径：国家流感中心北方省份真实周度 ILI% 监测序列，多源特征包含流感、气象与百度搜索指数。",
    ])


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("毕设答辩知识准备手册 | 基于深度学习和多元数据的流感爆发趋势预测")
        set_run_font(r)
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)


def main() -> None:
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_project_overview(doc)
    add_time_series_basics(doc)
    add_feature_engineering(doc)
    add_project_build_process(doc)
    add_itransformer(doc)
    add_baselines_metrics(doc)
    add_results(doc)
    add_qa(doc)
    add_memory_checklist(doc)
    add_references(doc)
    add_footer(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT)
        print(OUT)
    except PermissionError:
        doc.save(FALLBACK_OUT)
        print(FALLBACK_OUT)


if __name__ == "__main__":
    main()
